from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT, TA_RIGHT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import (
    BaseDocTemplate,
    Frame,
    KeepTogether,
    PageBreak,
    PageTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(__file__).resolve().parents[1]
PUBLIC_DIR = ROOT / "public" / "resume"
OUTPUT_DIR = ROOT / "deliverables"
PUBLIC_DIR.mkdir(parents=True, exist_ok=True)
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

PURPLE = "#6354E8"
INK = "#191A17"
MUTED = "#62655E"
LIGHT = "#E3E1DA"


PROFILE = (
    "Full-stack engineer and founder with 5+ years taking products from early ideas to production. "
    "Works across TypeScript, Python, React/Next.js, Node.js, data and infrastructure, with production "
    "experience in agentic workflows, RAG, document intelligence and realtime voice. Comfortable in "
    "zero-to-one environments that require product judgment, system architecture and hands-on execution."
)

EXPERIENCE_PAGE_1 = [
    {
        "role": "Founder & CTO",
        "company": "DcernX",
        "date": "2025 — Present",
        "bullets": [
            "Building an AI operating system for venture capital, private equity and investment teams in partnership with a UK-based accelerator, centered on agentic research, document intelligence and collaborative diligence.",
            "Designed multi-agent research workflows that coordinate web search, uploaded documents, specialized agents, tool calls and structured outputs through stateful orchestration.",
            "Built evidence-backed analysis with citation grounding, claim-to-evidence mapping, risk extraction and multi-document reasoning.",
            "Developed ingestion and retrieval pipelines using Docling, PageIndex, RAG and hierarchical/vector retrieval strategies.",
            "Own product and technical architecture across the Next.js/TypeScript application, Python AI workflows, data infrastructure, model/provider routing and deployment.",
        ],
    },
    {
        "role": "Founder / Product & Engineering Lead",
        "company": "P101 Labs",
        "date": "2024 — Present",
        "bullets": [
            "Built and shipped AI and full-stack products for startups and SMEs across voice, chat, interview automation, document intelligence and workflow automation.",
            "Partnered directly with founders to translate ambiguous operational problems into product scope, architecture and production-ready software.",
            "Owned interfaces, APIs, data models, integrations, streaming and asynchronous workflows alongside the AI systems embedded in each product.",
        ],
    },
    {
        "role": "Independent Consultant — Full-Stack & AI",
        "company": "Selected engagements",
        "date": "2023 — Present · Concurrent",
        "bullets": [
            "Partnered with startups and SMEs on product architecture, rapid prototyping and production engineering across operations, customer support and document intelligence.",
            "Contributed across product definition, frontend, backend, integrations and deployment based on the needs of each engagement.",
        ],
    },
]

SELECTED_PROJECTS = [
    {
        "title": "3D Garden & Home Designer",
        "label": "INTERACTIVE DESIGN PLATFORM",
        "summary": "A visual product that helps users create polished 3D garden and home mockups through an approachable design workflow.",
        "detail": "Worked across the interactive product experience and supporting platform, helping translate specialist 3D-design concepts into a tool for everyday users.",
    },
    {
        "title": "Krooki Management System",
        "label": "GOVTECH · WORKFLOW AUTOMATION",
        "summary": "A management platform supporting the Government of Oman in digitizing construction application and approval processes.",
        "detail": "Helped turn a multi-stage operational process into a structured digital workflow for application management, reviews, approvals and stakeholder coordination.",
    },
]

EXPERIENCE_PAGE_2 = [
    {
        "role": "Founding Full-Stack Engineer",
        "company": "CultureCo",
        "date": "Sep 2024 — Dec 2024",
        "bullets": [
            "Built the full-stack creator commerce platform enabling creators to sell products, manage communities and grow their businesses.",
            "Contributed across product development, frontend, backend and integrations in an early-stage, zero-to-one environment.",
        ],
    },
    {
        "role": "Founding Engineer",
        "company": "Truts",
        "date": "Dec 2021 — Mar 2024",
        "bullets": [
            "Built a Web3 discovery and quest platform from the ground up across frontend, backend, blockchain integrations and product development.",
            "Owned broad engineering scope as an early team member, taking product capabilities from concept through production release.",
        ],
    },
]

AI_SYSTEMS = [
    ("Deep research & diligence", "Multi-agent research across documents, web sources and structured data, producing grounded findings, citations, claim/evidence links and risk analysis."),
    ("Realtime voice agents", "Streaming STT → LLM → TTS systems with VAD, tool execution, LiveKit Agents, Whisper and production integration workflows."),
    ("AI interview systems", "Interview and screening products spanning question generation, response analysis, scoring, evaluation and video-introduction assessment."),
    ("Chatbots & workflow agents", "Knowledge-grounded, tool-enabled experiences with multi-turn state, external actions, configurable workflows and customer-support automation."),
]

STACK = [
    ("Product & frontend", "TypeScript · React · Next.js · Tailwind · React Query · Zustand · interaction design"),
    ("Backend & realtime", "Python · Node.js · Express · REST APIs · WebSockets/SSE · async workflows · system design"),
    ("Applied AI", "LangGraph · Vercel AI SDK · LangChain · LiveKit Agents · RAG · tool calling · structured outputs"),
    ("Models", "OpenAI · Claude · Gemini · DeepSeek · Qwen · Fireworks AI · open-weight inference · provider routing"),
    ("Data & infrastructure", "PostgreSQL · MongoDB · Convex · Supabase · Redis · AWS · Docker · Cloudflare · Railway · Coolify · Linux"),
]


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=90, bottom=80, end=90):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_docx_rule(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "5")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "D8D6CF")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def docx_section_heading(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(title.upper())
    run.bold = True
    run.font.name = "Aptos"
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(99, 84, 232)
    add_docx_rule(p)


def docx_experience(doc, item):
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Inches(5.3)
    table.columns[1].width = Inches(1.6)
    left, right = table.rows[0].cells
    for cell in (left, right):
        set_cell_margins(cell, top=40, start=0, bottom=40, end=0)
    p = left.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    role = p.add_run(item["role"])
    role.bold = True
    role.font.size = Pt(10.7)
    role.font.color.rgb = RGBColor(25, 26, 23)
    company = p.add_run(f"  ·  {item['company']}")
    company.font.size = Pt(10)
    company.font.color.rgb = RGBColor(74, 76, 70)
    p = right.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(item["date"])
    run.font.size = Pt(8.4)
    run.font.color.rgb = RGBColor(98, 101, 94)
    for bullet in item["bullets"]:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.18)
        p.paragraph_format.first_line_indent = Inches(-0.12)
        p.paragraph_format.space_after = Pt(3.2)
        p.paragraph_format.line_spacing = 1.08
        run = p.add_run(bullet)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(55, 57, 52)


def build_docx(path):
    doc = Document()
    section = doc.sections[0]
    section.page_height = Inches(11.69)
    section.page_width = Inches(8.27)
    section.top_margin = Inches(0.62)
    section.bottom_margin = Inches(0.62)
    section.left_margin = Inches(0.66)
    section.right_margin = Inches(0.66)

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(9.2)
    normal.font.color.rgb = RGBColor(25, 26, 23)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("JOHN SWAROOP")
    run.bold = True
    run.font.name = "Georgia"
    run.font.size = Pt(29)
    run.font.color.rgb = RGBColor(25, 26, 23)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(9)
    run = p.add_run("Founder · Full-Stack Engineer · Applied AI")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(99, 84, 232)

    contact = doc.add_table(rows=1, cols=2)
    contact.autofit = False
    c1, c2 = contact.rows[0].cells
    for cell in (c1, c2):
        set_cell_shading(cell, "F0EEE8")
        set_cell_margins(cell, top=120, start=130, bottom=120, end=130)
    c1.text = "Hyderabad, India  ·  +91 9618788956  ·  johnswaroop28@gmail.com"
    c2.text = "github.com/johnswaroop  ·  linkedin.com/in/john-swaroop-4389961b7  ·  dcernx.com"
    c2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for cell in (c1, c2):
        for run in cell.paragraphs[0].runs:
            run.font.size = Pt(7.6)
            run.font.color.rgb = RGBColor(75, 77, 72)

    docx_section_heading(doc, "Profile")
    p = doc.add_paragraph(PROFILE)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.14
    p.runs[0].font.size = Pt(9.5)

    docx_section_heading(doc, "Core strengths")
    strengths = doc.add_table(rows=1, cols=3)
    strengths.autofit = False
    strength_data = [
        ("ZERO-TO-ONE PRODUCT", "Product strategy · architecture · prototyping · technical leadership · launch"),
        ("FULL-STACK ENGINEERING", "TypeScript · Python · React/Next.js · Node.js · APIs · data · infrastructure"),
        ("APPLIED AI SYSTEMS", "Agents · RAG · voice · document intelligence · tool calling · model routing"),
    ]
    for cell, (title, detail) in zip(strengths.rows[0].cells, strength_data):
        set_cell_shading(cell, "F4F2ED")
        set_cell_margins(cell, top=140, start=120, bottom=140, end=120)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(title)
        r.bold = True
        r.font.size = Pt(7.8)
        r.font.color.rgb = RGBColor(99, 84, 232)
        p = cell.add_paragraph(detail)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.05
        p.runs[0].font.size = Pt(8.2)

    docx_section_heading(doc, "Professional experience")
    for item in EXPERIENCE_PAGE_1:
        docx_experience(doc, item)

    doc.add_page_break()
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("SELECTED WORK & TECHNICAL DEPTH")
    r.font.name = "Georgia"
    r.font.size = Pt(20)
    r.font.color.rgb = RGBColor(25, 26, 23)
    p = doc.add_paragraph("A closer look at selected product work and the systems behind it.")
    p.paragraph_format.space_after = Pt(4)
    p.runs[0].font.size = Pt(9)
    p.runs[0].font.color.rgb = RGBColor(98, 101, 94)

    docx_section_heading(doc, "Selected product work")
    projects = doc.add_table(rows=1, cols=2)
    for cell, project in zip(projects.rows[0].cells, SELECTED_PROJECTS):
        set_cell_shading(cell, "F4F2ED")
        set_cell_margins(cell, top=150, start=150, bottom=150, end=150)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(5)
        r = p.add_run(project["label"])
        r.bold = True
        r.font.size = Pt(7.4)
        r.font.color.rgb = RGBColor(99, 84, 232)
        p = cell.add_paragraph(project["title"])
        p.paragraph_format.space_after = Pt(6)
        r = p.runs[0]
        r.bold = True
        r.font.name = "Georgia"
        r.font.size = Pt(12)
        p = cell.add_paragraph(project["summary"])
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.line_spacing = 1.08
        p.runs[0].font.size = Pt(8.7)
        p = cell.add_paragraph(project["detail"])
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.08
        p.runs[0].font.size = Pt(8.3)
        p.runs[0].font.color.rgb = RGBColor(83, 85, 79)

    docx_section_heading(doc, "Earlier experience")
    for item in EXPERIENCE_PAGE_2:
        docx_experience(doc, item)

    docx_section_heading(doc, "Selected applied AI systems")
    systems = doc.add_table(rows=2, cols=2)
    for cell, (title, detail) in zip([c for row in systems.rows for c in row.cells], AI_SYSTEMS):
        set_cell_margins(cell, top=110, start=120, bottom=110, end=120)
        set_cell_shading(cell, "F4F2ED")
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(title)
        r.bold = True
        r.font.size = Pt(8.8)
        p = cell.add_paragraph(detail)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.08
        p.runs[0].font.size = Pt(8.1)
        p.runs[0].font.color.rgb = RGBColor(83, 85, 79)

    docx_section_heading(doc, "Technical stack")
    for label, detail in STACK:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4.5)
        r = p.add_run(f"{label}: ")
        r.bold = True
        r.font.size = Pt(8.4)
        r = p.add_run(detail)
        r.font.size = Pt(8.4)
        r.font.color.rgb = RGBColor(72, 74, 69)

    doc.save(path)


def pdf_styles():
    return {
        "name": ParagraphStyle("Name", fontName="Times-Roman", fontSize=30, leading=31, textColor=colors.HexColor(INK), spaceAfter=3),
        "title": ParagraphStyle("Title", fontName="Helvetica", fontSize=10.5, leading=14, textColor=colors.HexColor(PURPLE)),
        "page_title": ParagraphStyle("PageTitle", fontName="Times-Roman", fontSize=20, leading=23, textColor=colors.HexColor(INK), spaceAfter=4),
        "page_intro": ParagraphStyle("PageIntro", fontName="Helvetica", fontSize=8.7, leading=12, textColor=colors.HexColor(MUTED), spaceAfter=5),
        "contact": ParagraphStyle("Contact", fontName="Helvetica", fontSize=7.3, leading=10, textColor=colors.HexColor(MUTED)),
        "section": ParagraphStyle("Section", fontName="Helvetica-Bold", fontSize=8.1, leading=11, tracking=1.25, textColor=colors.HexColor(PURPLE), spaceBefore=0, spaceAfter=0),
        "body": ParagraphStyle("Body", fontName="Helvetica", fontSize=9.1, leading=13.1, textColor=colors.HexColor("#373934")),
        "role": ParagraphStyle("Role", fontName="Helvetica-Bold", fontSize=10.2, leading=13, textColor=colors.HexColor(INK)),
        "date": ParagraphStyle("Date", fontName="Helvetica", fontSize=8, leading=11, alignment=TA_RIGHT, textColor=colors.HexColor(MUTED)),
        "bullet": ParagraphStyle("Bullet", fontName="Helvetica", fontSize=8.65, leading=11.7, leftIndent=11, firstLineIndent=-7, bulletIndent=0, textColor=colors.HexColor("#373934"), spaceAfter=3),
        "card_title": ParagraphStyle("CardTitle", fontName="Helvetica-Bold", fontSize=8, leading=10, textColor=colors.HexColor(PURPLE), spaceAfter=4),
        "card_body": ParagraphStyle("CardBody", fontName="Helvetica", fontSize=8.7, leading=12, textColor=colors.HexColor("#4D4F49")),
        "project_label": ParagraphStyle("ProjectLabel", fontName="Helvetica-Bold", fontSize=7.2, leading=9, tracking=.8, textColor=colors.HexColor(PURPLE), spaceAfter=5),
        "project_title": ParagraphStyle("ProjectTitle", fontName="Times-Roman", fontSize=15, leading=17, textColor=colors.HexColor(INK), spaceAfter=7),
        "project_body": ParagraphStyle("ProjectBody", fontName="Helvetica", fontSize=9, leading=12.5, textColor=colors.HexColor("#3E403A"), spaceAfter=7),
        "project_detail": ParagraphStyle("ProjectDetail", fontName="Helvetica", fontSize=8.5, leading=12, textColor=colors.HexColor(MUTED)),
        "stack_label": ParagraphStyle("StackLabel", fontName="Helvetica-Bold", fontSize=7.8, leading=10, textColor=colors.HexColor(PURPLE)),
        "stack": ParagraphStyle("Stack", fontName="Helvetica", fontSize=8.5, leading=11.2, textColor=colors.HexColor("#454741")),
    }


def section_title(text, styles):
    content = Paragraph(text.upper(), styles["section"])
    table = Table([[content]], colWidths=[184 * mm])
    table.spaceBefore = 12
    table.spaceAfter = 7
    table.keepWithNext = True
    table.setStyle(TableStyle([
        ("LINEBELOW", (0, 0), (-1, -1), 0.45, colors.HexColor(LIGHT)),
        ("LEFTPADDING", (0, 0), (-1, -1), 0),
        ("RIGHTPADDING", (0, 0), (-1, -1), 0),
        ("TOPPADDING", (0, 0), (-1, -1), 0),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    return table


def experience_block(item, styles):
    heading = Table([[
        Paragraph(f"{item['role']} <font color='#62655E'> · {item['company']}</font>", styles["role"]),
        Paragraph(item["date"], styles["date"]),
    ]], colWidths=[139 * mm, 45 * mm])
    heading.setStyle(TableStyle([
        ("LEFTPADDING", (0, 0), (-1, -1), 0), ("RIGHTPADDING", (0, 0), (-1, -1), 0),
        ("TOPPADDING", (0, 0), (-1, -1), 1), ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
    ]))
    bullets = [Paragraph(f"• {bullet}", styles["bullet"]) for bullet in item["bullets"]]
    return KeepTogether([heading, *bullets, Spacer(1, 7)])


def pdf_header_footer(canvas, doc):
    canvas.saveState()
    canvas.setTitle("John Swaroop — Founder, Full-Stack Engineer & Applied AI")
    canvas.setAuthor("John Swaroop")
    canvas.setStrokeColor(colors.HexColor(PURPLE))
    canvas.setLineWidth(1.2)
    canvas.line(13 * mm, 12 * mm, 31 * mm, 12 * mm)
    canvas.setFont("Helvetica", 6.5)
    canvas.setFillColor(colors.HexColor("#85877F"))
    canvas.drawString(34 * mm, 10.3 * mm, "JOHN SWAROOP  ·  FOUNDER / FULL-STACK / APPLIED AI")
    canvas.drawRightString(197 * mm, 10.3 * mm, f"{doc.page} / 2")
    canvas.restoreState()


def build_pdf(path):
    styles = pdf_styles()
    doc = BaseDocTemplate(str(path), pagesize=A4, leftMargin=13 * mm, rightMargin=13 * mm, topMargin=13 * mm, bottomMargin=16 * mm)
    frame = Frame(doc.leftMargin, doc.bottomMargin, doc.width, doc.height, id="normal")
    doc.addPageTemplates([PageTemplate(id="main", frames=frame, onPage=pdf_header_footer)])
    story = []
    story.append(Paragraph("JOHN SWAROOP", styles["name"]))
    story.append(Paragraph("Founder · Full-Stack Engineer · Applied AI", styles["title"]))
    story.append(Spacer(1, 8))
    contact = Table([[
        Paragraph("Hyderabad, India · +91 9618788956 · <link href='mailto:johnswaroop28@gmail.com'>johnswaroop28@gmail.com</link>", styles["contact"]),
        Paragraph("<link href='https://github.com/johnswaroop'>GitHub</link> · <link href='https://linkedin.com/in/john-swaroop-4389961b7'>LinkedIn</link> · <link href='https://dcernx.com'>dcernx.com</link>", ParagraphStyle("ContactRight", parent=styles["contact"], alignment=TA_RIGHT)),
    ]], colWidths=[103 * mm, 81 * mm])
    contact.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#F0EEE8")),
        ("LEFTPADDING", (0, 0), (-1, -1), 9), ("RIGHTPADDING", (0, 0), (-1, -1), 9),
        ("TOPPADDING", (0, 0), (-1, -1), 8), ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
    ]))
    story.append(contact)
    story.append(section_title("Profile", styles))
    story.append(Paragraph(PROFILE, styles["body"]))
    story.append(section_title("Core strengths", styles))
    strength_cells = []
    for title, detail in [
        ("ZERO-TO-ONE PRODUCT", "Product strategy · architecture · prototyping · technical leadership · launch"),
        ("FULL-STACK ENGINEERING", "TypeScript · Python · React/Next.js · Node.js · APIs · data · infrastructure"),
        ("APPLIED AI SYSTEMS", "Agents · RAG · voice · document intelligence · tool calling · model routing"),
    ]:
        strength_cells.append([Paragraph(title, styles["card_title"]), Paragraph(detail, styles["card_body"])])
    strength_table = Table([[cell for cell in strength_cells]], colWidths=[61.33 * mm] * 3)
    strength_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#F4F2ED")),
        ("LINEABOVE", (0, 0), (-1, -1), 1.1, colors.HexColor(PURPLE)),
        ("LEFTPADDING", (0, 0), (-1, -1), 9), ("RIGHTPADDING", (0, 0), (-1, -1), 9),
        ("TOPPADDING", (0, 0), (-1, -1), 9), ("BOTTOMPADDING", (0, 0), (-1, -1), 9),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
    ]))
    story.append(strength_table)
    story.append(section_title("Professional experience", styles))
    for item in EXPERIENCE_PAGE_1:
        story.append(experience_block(item, styles))

    story.append(PageBreak())
    story.append(Paragraph("SELECTED WORK &amp; TECHNICAL DEPTH", styles["page_title"]))
    story.append(Paragraph("A closer look at selected product work and the systems behind it.", styles["page_intro"]))

    story.append(section_title("Selected product work", styles))
    project_cells = []
    for project in SELECTED_PROJECTS:
        project_cells.append([
            Paragraph(project["label"], styles["project_label"]),
            Paragraph(project["title"], styles["project_title"]),
            Paragraph(project["summary"], styles["project_body"]),
            Paragraph(project["detail"], styles["project_detail"]),
        ])
    projects_table = Table([project_cells], colWidths=[90.5 * mm, 90.5 * mm], hAlign="LEFT")
    projects_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#F4F2ED")),
        ("LINEABOVE", (0, 0), (-1, -1), 1.2, colors.HexColor(PURPLE)),
        ("BOX", (0, 0), (-1, -1), 0.45, colors.HexColor(LIGHT)),
        ("INNERGRID", (0, 0), (-1, -1), 0.45, colors.HexColor(LIGHT)),
        ("LEFTPADDING", (0, 0), (-1, -1), 13), ("RIGHTPADDING", (0, 0), (-1, -1), 13),
        ("TOPPADDING", (0, 0), (-1, -1), 15), ("BOTTOMPADDING", (0, 0), (-1, -1), 15),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
    ]))
    story.append(projects_table)

    story.append(section_title("Earlier experience", styles))
    for item in EXPERIENCE_PAGE_2:
        story.append(experience_block(item, styles))

    story.append(section_title("Selected applied AI systems", styles))
    system_cells = []
    for title, detail in AI_SYSTEMS:
        system_cells.append([Paragraph(title, styles["role"]), Paragraph(detail, styles["card_body"])])
    systems_table = Table([system_cells[:2], system_cells[2:]], colWidths=[92 * mm, 92 * mm])
    systems_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#F4F2ED")),
        ("BOX", (0, 0), (-1, -1), 0.4, colors.HexColor(LIGHT)),
        ("INNERGRID", (0, 0), (-1, -1), 0.4, colors.HexColor(LIGHT)),
        ("LEFTPADDING", (0, 0), (-1, -1), 11), ("RIGHTPADDING", (0, 0), (-1, -1), 11),
        ("TOPPADDING", (0, 0), (-1, -1), 12), ("BOTTOMPADDING", (0, 0), (-1, -1), 12),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
    ]))
    story.append(systems_table)

    story.append(section_title("Technical stack", styles))
    stack_rows = [[Paragraph(label.upper(), styles["stack_label"]), Paragraph(detail, styles["stack"])] for label, detail in STACK]
    stack_table = Table(stack_rows, colWidths=[38 * mm, 146 * mm], hAlign="LEFT")
    stack_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), colors.white),
        ("BACKGROUND", (0, 1), (-1, 1), colors.HexColor("#F7F5F0")),
        ("BACKGROUND", (0, 3), (-1, 3), colors.HexColor("#F7F5F0")),
        ("LINEBELOW", (0, 0), (-1, -1), 0.35, colors.HexColor(LIGHT)),
        ("LEFTPADDING", (0, 0), (-1, -1), 7), ("RIGHTPADDING", (0, 0), (-1, -1), 7),
        ("TOPPADDING", (0, 0), (-1, -1), 7), ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
    ]))
    story.append(stack_table)
    doc.build(story)


if __name__ == "__main__":
    docx_path = OUTPUT_DIR / "John-Swaroop-Resume.docx"
    pdf_path = OUTPUT_DIR / "John-Swaroop-Resume.pdf"
    build_docx(docx_path)
    build_pdf(pdf_path)
    (PUBLIC_DIR / pdf_path.name).write_bytes(pdf_path.read_bytes())
    print(docx_path)
    print(pdf_path)
    print(PUBLIC_DIR / pdf_path.name)
